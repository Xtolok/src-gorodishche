"""Генерация Word-таблицы файлов проекта."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROWS = [
    ("Корень проекта", [
        ("manage.py", "manage.py", "Точка входа Django: сервер, миграции, management-команды"),
        ("requirements.txt", "requirements.txt", "Список зависимостей: Django и Pillow"),
        (".gitignore", ".gitignore", "Исключения для Git: venv, БД, кэш, IDE"),
        ("db.sqlite3", "db.sqlite3", "Файл базы данных SQLite (контент сайта)"),
    ]),
    ("Пакет config/ (настройки проекта)", [
        ("config/__init__.py", "__init__.py", "Пустой маркер Python-пакета"),
        ("config/settings.py", "settings.py", "Настройки Django: приложения, БД, шаблоны, статика, медиа"),
        ("config/urls.py", "urls.py", "Корневые URL: подключение main, админка, раздача media в DEBUG"),
        ("config/wsgi.py", "wsgi.py", "WSGI-приложение для деплоя (Apache, Gunicorn и т.п.)"),
        ("config/asgi.py", "asgi.py", "ASGI-приложение для асинхронного деплоя"),
    ]),
    ("Приложение main/ (логика сайта)", [
        ("main/__init__.py", "__init__.py", "Пустой маркер Python-пакета приложения"),
        ("main/apps.py", "apps.py", "Регистрация приложения main в Django"),
        ("main/models.py", "models.py", "Модели БД: страницы, новости, услуги, контакты, настройки"),
        ("main/views.py", "views.py", "Обработчики всех публичных страниц сайта"),
        ("main/urls.py", "urls.py", "Маршруты страниц: главная, услуги, контакты, новости и др."),
        ("main/admin.py", "admin.py", "Админ-панель для редактирования контента сайта"),
        ("main/context_processors.py", "context_processors.py", "Общие данные в шаблонах: название сайта и контакты"),
        ("main/tests.py", "tests.py", "Заготовка для автотестов (пока пустая)"),
        ("main/management/__init__.py", "__init__.py", "Маркер пакета management-команд"),
        ("main/management/commands/__init__.py", "__init__.py", "Маркер пакета пользовательских команд"),
        (
            "main/management/commands/populate_site.py",
            "populate_site.py",
            "Команда populate_site: начальное заполнение БД демо-данными",
        ),
        ("main/migrations/__init__.py", "__init__.py", "Маркер пакета миграций"),
        ("main/migrations/0001_initial.py", "0001_initial.py", "Первая миграция: создание всех таблиц приложения"),
        (
            "main/migrations/0002_contactinfo_map_iframe_contactinfo_map_latitude_and_more.py",
            "0002_contactinfo_map_iframe_contactinfo_map_latitude_and_more.py",
            "Вторая миграция: поля карты (координаты, iframe, zoom)",
        ),
    ]),
    ("Шаблоны templates/", [
        ("templates/base.html", "base.html", "Базовый каркас HTML: шапка, подвал, подключение CSS"),
        ("templates/includes/header.html", "header.html", "Шапка сайта: логотип, меню, ссылка на Госуслуги"),
        ("templates/includes/footer.html", "footer.html", "Подвал сайта: копирайт и служебные ссылки"),
        ("templates/pages/index.html", "index.html", "Вёрстка главной: статистика, карточки, блок «как получить»"),
        ("templates/pages/socialnye-uslugi.html", "socialnye-uslugi.html", "Страница «Социальные услуги»"),
        ("templates/pages/perechen-uslug.html", "perechen-uslug.html", "Страница «Перечень услуг» по категориям"),
        ("templates/pages/dnevnoe-otdelenie.html", "dnevnoe-otdelenie.html", "Страница «Дневное отделение»"),
        ("templates/pages/rabota-s-obrascheniyami.html", "rabota-s-obrascheniyami.html", "Страница «Работа с обращениями»"),
        ("templates/pages/kontakty.html", "kontakty.html", "Страница «Контакты»: адрес, телефоны, карта, приём"),
        ("templates/pages/novosti.html", "novosti.html", "Страница «Новости»: список публикаций"),
        ("templates/pages/pravovye-akty.html", "pravovye-akty.html", "Страница «Правовые акты»: разделы и ссылки на документы"),
    ]),
    ("Статика и медиа", [
        ("static/css/site.css", "site.css", "Стили оформления всего сайта"),
        ("media/news/images.jpg", "images.jpg", "Загруженное изображение к новости"),
        ("media/news/images_omaC1CD.jpg", "images_omaC1CD.jpg", "Загруженное изображение к новости (копия с уникальным именем)"),
        ("media/news/images_qiTqCrx.jpg", "images_qiTqCrx.jpg", "Загруженное изображение к новости (копия с уникальным именем)"),
    ]),
]

HEADERS = ("Обозначение файла (путь)", "Наименование файла", "Примечание")


def add_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    for col, text in enumerate(HEADERS):
        cell = table.rows[0].cells[col]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, (path, name, note) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = path
        table.rows[i].cells[1].text = name
        table.rows[i].cells[2].text = note
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
    widths = (Cm(7), Cm(5), Cm(8))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def main():
    out = Path(__file__).resolve().parent.parent / "Описание_файлов_проекта.docx"
    doc = Document()
    title = doc.add_heading("Описание файлов проекта", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Сайт Городищенского СРЦ (Django). Всего 35 файлов (без .venv и __pycache__)."
    )
    for section_title, rows in ROWS:
        doc.add_heading(section_title, level=2)
        add_table(doc, rows)
        doc.add_paragraph()
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
